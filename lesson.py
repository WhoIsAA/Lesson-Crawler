#! /usr/bin/env python3
from urllib import request
from urllib import error
from urllib import parse
from http import cookiejar
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime
import os


'''
根据范围截取字符串
'''
def analyse(html, start_s, end_s):
	start = html.find(start_s) + len(start_s)
	tmp_html = html[start:]
	end = tmp_html.find(end_s)
	return tmp_html[:end].strip()

'''
解析课程列表
'''
def analyse_lesson(opener, headers, html):
	#获取课程名称
	tmp_folder = analyse(html, "<div class=\"database-title clearfix\">", "</div>")
	folder = analyse(tmp_folder, "<span>", "</span>")
	#创建文件夹，改变当前工作目录
	print("正在创建文件夹（%s）..." % folder)
	if not os.path.exists(folder):
		os.mkdir(folder)
	os.chdir(folder)

	#循环获取每一个课程的试题
	lesson_html = analyse(html, "<ul class=\"lesson-chap-ul\">", "</ul>")
	while True:
		tmp_html = analyse(lesson_html, "<li class=\"clearfix\">", "</li>")
		lesson_html = analyse(lesson_html, tmp_html, "</ul>")
		sectionid = analyse(tmp_html, "index.php/Lessontiku/questionsmore_manage/sectionid/", "/subjectid")
		analyse_exam(opener, headers, tmp_html, sectionid)

		if not tmp_html or not lesson_html:
			break;

'''
解析试题标题、题目总数
'''
def analyse_exam(opener, headers, html, sectionid):
	#获取标题
	title = analyse(html, "<div class=\"lesson-errchap-tit\">", "</div>")
	#获取题目总数
	total_size = analyse(html, "<span class=\"progressNum\">", "</span>")
	start = total_size.find("/") + 1
	total_size = total_size[start:]
	print("正在下载（%s） 题目总数：%s" % (title, total_size))

	#考题，添加标题
	exam_doc = Document()
	heading = exam_doc.add_heading(title, 0)
	heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

	#答案，添加标题
	answers_doc = Document()
	heading = answers_doc.add_heading(title + "（答案）", 0)
	heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

	#循环解析题目
	for index in range(1, int(total_size) + 1):
		result_url = 'http://i.sxmaps.com/index.php/Lessontiku/questionsmore_manage/subjectid/111/sectionid/%s/p/%d/majorid_sx/38/classid_sx/24' % (sectionid, index)
		item_request = request.Request(result_url, headers = headers)
		try:
			response = opener.open(item_request)
			html = response.read().decode('utf-8')
			exam_doc = analyse_item(index, html, exam_doc)
			answers_doc = analyse_answers(index, html, answers_doc)
		except error.URLError as e:
			if hasattr(e, 'code'):
				print("HTTPError:%d" % e.code)
			elif hasattr(e, 'reason'):
				print("URLError:%s" % e.reason)

	filename = "%s.docx" % title
	exam_doc.save(filename)
	print("成功创建文件：%s" % filename)
	filename = "%s（答案）.docx" % title
	answers_doc.save(filename)
	print("成功创建文件：%s" % filename)

'''
解析每道试题详细信息
'''
def analyse_item(index, html, document):
	#题目类型
	type_s = "<div class=\"database-txt\">"
	start = html.find(type_s) + len(type_s)
	tmp_html = html[start:]
	end = tmp_html.find("</a>")
	start = end - 5
	exam_type = tmp_html[start:end].strip()

	#标题
	title = analyse(tmp_html, "<pre>", "</pre>")
	paragraph = "%d.%s %s" % (index, exam_type, title)
	document.add_paragraph(paragraph)
	print("标题：%s" % paragraph)

	if(exam_type == '[单选题]' or exam_type == '[多选题]'):
		#选项
		options = []
		while True:
			option_s = "<div class=\"lesson-xz-txt\">"
			end_s = "<div class=\"hide\" onclick=\"lesson.isQuestionJxShow()\">确定</div>"
			end_div_s = "</div>"

			if tmp_html.find(option_s) <= 0:
				break

			start = tmp_html.find(option_s) + len(option_s)
			end = tmp_html.find(end_s)
			tmp_html = tmp_html[start:end]
			end = tmp_html.find(end_div_s)
			option = tmp_html[:end].strip()
			document.add_paragraph(option)
			options.append(option)
		print("选项：%s" % options)
	elif(exam_type == '简答题'):
		document.add_paragraph("")
		document.add_paragraph("")
		document.add_paragraph("")
	#加入一个空白行
	document.add_paragraph("")
	return document

'''
解析每道试题的正确答案
'''
def analyse_answers(index, html, document):
	#正确答案
	right_s = "<pre style='line-height: 1.5;white-space: pre-wrap;'>"
	right = "%s.正确答案：%s" % (index, analyse(html, right_s, "</pre>"))
	print(right)
	document.add_paragraph(right)
	return document


if __name__ == '__main__':
	list_url = "http://i.sxmaps.com/index.php/lessontiku/questions_manage/subjectid/111/classid_sx/24/majorid_sx/38.html"
	
	headers = {
		'User-Agent': 'Mozilla/5.0 (Windows NT 6.2; WOW64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/27.0.1453.94 Safari/537.36',
		'Connection': 'keep-alive', 
		'Cookie': '【请先使用Fiddler抓取并输入你的Cookie】'
	}

	cookie = cookiejar.CookieJar()
	handler = request.HTTPCookieProcessor(cookie)
	opener = request.build_opener(handler)

	list_request = request.Request(list_url, headers = headers)
	try:
		response = opener.open(list_request)
		html = response.read().decode('utf-8')
		start_t = datetime.now()
		analyse_lesson(opener, headers, html)
		end_t = datetime.now()
		print("*" * 80)
		print("* 下载完成，总共用了%s秒。" % (end_t - start_t).seconds)
		print("*" * 80)
	except error.URLError as e:
		if hasattr(e, 'code'):
			print("HTTPError:%d" % e.code)
		elif hasattr(e, 'reason'):
			print("URLError:%s" % e.reason)


