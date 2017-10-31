#! /usr/bin/env python3
from urllib import request
from urllib import error
from urllib import parse
from http import cookiejar
from bs4 import BeautifulSoup
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime
import os


'''
去掉空格、换行符
'''
def strip(str):
    if str:
        return str.replace('\r','').replace('\n','').replace('\t','').strip()
    else:
        return ''

'''
根据范围截取字符串
'''
def intercept_string(str, start_s, end_s):
    start = str.find(start_s) + len(start_s)
    tmp = str[start:]
    end = tmp.find(end_s)
    return tmp[:end].strip()

'''
解析课程列表
'''
def analyse_lesson(opener, headers, html):
    soup = BeautifulSoup(html, 'lxml')
    #获取课程名称
    folder = strip(soup.find('div', {'class': 'database-title clearfix'}).text)
    #创建文件夹，改变当前工作目录
    print("正在创建文件夹（%s）..." % folder)
    if not os.path.exists(folder):
        os.mkdir(folder)
    os.chdir(folder)

    for li in soup.find_all('li', {'class': 'clearfix'}):
        title = li.find('div', {'class': 'lesson-errchap-tit'}).text;
        count = li.find('span', {'class': 'progressNum'}).text.split("/")[1]
        section_id = li.find('div', {'class': 'lesson-re-do'})['onclick']
        section_id = intercept_string(section_id, "sectionid/", "/subjectid")
        print("正在下载（%s） 题目总数：%s" % (title, count))
        analyse_exam(opener, headers, title, count, section_id)

'''
解析试题标题、题目总数
'''
def analyse_exam(opener, headers, title, count, sectionid):
    #考题，添加标题
    exam_doc = Document()
    heading = exam_doc.add_heading(title, 0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #答案，添加标题
    answers_doc = Document()
    heading = answers_doc.add_heading(title + "（答案）", 0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #循环解析题目
    for index in range(1, int(count) + 1):
        result_url = 'http://i.sxmaps.com/index.php/Lessontiku/questionsmore_manage/subjectid/111/sectionid/%s/p/%d/majorid_sx/38/classid_sx/24' % (sectionid, index)
        item_request = request.Request(result_url, headers = headers)
        try:
            response = opener.open(item_request)
            html = response.read().decode('utf-8')
            soup = BeautifulSoup(html, 'lxml')
            exam_doc = analyse_item(soup, exam_doc)
            answers_doc = analyse_answers(index, soup, answers_doc)
        except error.URLError as e:
            if hasattr(e, 'code'):
                print("HTTPError:%d" % e.code)
            elif hasattr(e, 'reason'):
                print("URLError:%s" % e.reason)

    filename = "%s.docx" % title
    exam_doc.save(filename)
    print("成功创建文件：%s" % filename)
    filename = "%s（答案）.docx" % title
    answers_doc.save(filename)
    print("成功创建文件：%s" % filename)

'''
解析每道试题详细信息
'''
def analyse_item(soup, document):
    exam = soup.find('div', {'class': 'database-txt'})
    #题目类型
    exam_type = strip(exam.a.text)
    #题目序号
    exam_index = strip(exam.em.text)
    #题目内容
    exam_con = strip(exam.pre.text)
    paragraph = "%s%s %s" % (exam_index, exam_type, exam_con)
    document.add_paragraph(paragraph)
    print("标题：%s" % paragraph)

    if(exam_type == '[单选题]' or exam_type == '[多选题]'):
        #选项
        for option in soup.find_all('div', {'class': 'lesson-xz-txt'}):
            document.add_paragraph(strip(option.text))
    elif(exam_type == '[简答题]'):
        for i in range(3):
            document.add_paragraph("")
    #加入一个空白行
    document.add_paragraph("")
    return document

'''
解析每道试题的正确答案
'''
def analyse_answers(index, soup, document):
    #正确答案
    right = soup.find('div', {'class': 'lesson-da-desc'})
    right = "%s.正确答案：%s" % (index, strip(right.pre.text))
    print(right)
    document.add_paragraph(right)
    return document


if __name__ == '__main__':
	login_url = "http://i.sxmaps.com/index.php/member/login.html"
	list_url = "http://i.sxmaps.com/index.php/lessontiku/questions_manage/subjectid/111/classid_sx/24/majorid_sx/38.html"
	
	#请求头
	headers = {
		'User-Agent': 'Mozilla/5.0 (Windows NT 6.2; WOW64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/27.0.1453.94 Safari/537.36',
		'Connection': 'keep-alive', 
		'DNT': '1',
		'Referer': 'http://i.sxmaps.com/index.php/member/login.html',
		'Origin': 'http://i.sxmaps.com',
	}

	#请求参数
	data = {}
	data['password'] = "611640"
	data['phone'] = "15602611640"
	data['rember_me'] = "0"
	logingData = parse.urlencode(data).encode('utf-8')

	cookie = cookiejar.CookieJar()
	handler = request.HTTPCookieProcessor(cookie)
	opener = request.build_opener(handler)

    #登录请求
	login_request = request.Request(url=login_url, data=logingData, headers=headers)
	#课程列表请求
	list_request = request.Request(list_url, headers = headers)
	try:
    	#模拟登录
		login_rsp = opener.open(login_request)
		response = opener.open(list_request)
		html = response.read().decode('utf-8')
		start_t = datetime.now()
		analyse_lesson(opener, headers, html)
		end_t = datetime.now()
		print("*" * 80)
		print("* 下载完成，总共用了%s秒。" % (end_t - start_t).seconds)
		print("*" * 80)
	except error.URLError as e:
		if hasattr(e, 'code'):
			print("HTTPError:%d" % e.code)
		elif hasattr(e, 'reason'):
			print("URLError:%s" % e.reason)

